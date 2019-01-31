import urllib3
import urllib.request
from bs4 import BeautifulSoup
from post import Post
import os
from docx import Document
from docx.shared import Inches


#generates links across full range of dates
def make_links():
    """make full range of dates to use"""
    extensions = []
    for i in range(1, 9):
        for j in range(1, 13):
            if len(str(j)) == 1:
                extensions.append("/201{}/0{}".format(i, j))
            else:
                extensions.append("/201{}/{}".format(i, j))
    return extensions


#gets most blog posts
def main_scrape(url):
    """A function that scrapes the main page"""
    http = urllib3.PoolManager()
    req = http.request('GET', url)
    if req.status != 200:
        print("Error:", req.status, "skipping page", url)
        return None
    page = req.data
    soup = BeautifulSoup(page, 'lxml')
    contents = soup.find('div', {'id': 'primary-content'})
    articles = contents.find_all("div", {'class': "post-wrapper"})
    posts = []
    for art in articles:
        #title and date
        title = (art.find('h2', {'class': 'post-title'})).find('a').string
        date = (art.find('span', {"class": "entry-date"})).string
        post = Post(title, date)
        #text
        body = art.find("div", {"class": 'entry'})
        i = 0
        for paragraph in body.find_all('p'):
            if paragraph.find('em') is not None:
                if i == 0:
                    post.set_subtitle((paragraph.find('em')).string)
                else:
                    post.set_endnotes((paragraph.find('em')).string)
            else:
                post.add_paragraph(paragraph.string)
            i += 1
        #images
        try:
            for i in body.find_all('img'):
                try:
                    imgUrl = str(i['data-orig-file'])
                    imgName = imgUrl.split(".com")[1].replace("/", "_")
                    post.add_image([imgName, None]) # name, caption
                    #urllib.request.urlretrieve(imgUrl, os.path.basename("img/" + imgName))
                except KeyError:
                    imgUrl = str(i['src']).split('?')[0]
                    imgName = imgUrl.split(".com")[1].replace("/", "_")
                    post.add_image([imgName, None]) # name, caption
                    #urllib.request.urlretrieve(imgUrl, os.path.basename("img/" + imgName))
            for d in body.find_all('div', {"class": "wp-caption"}):
                try:
                    imgUrl = str(d.find('img')['data-orig-file'])
                    imgName = imgUrl.split(".com")[1].replace("/", "_")
                    post.add_image([imgName, (d.find('p')).string])
                    #urllib.request.urlretrieve(imgUrl, os.path.basename("img/" + imgName))
                except KeyError:
                    imgUrl = str(d.find('img')['src']).split('?')[0]
                    imgName = imgUrl.split(".com")[1].replace("/", "_")
                    post.add_image([imgName, (d.find('p')).string])
                    #urllib.request.urlretrieve(imgUrl, os.path.basename("img/" + imgName))
        except: #skip WP junk
            pass
        #end
        posts.append(post)
    for p in reversed(posts):
        global_articles.append(p)
    print("done with", url)


#gets a special post and inserts it correctly
def special_scrape(url):
    """a function that handles a single edge-case page"""
    http = urllib3.PoolManager()
    req = http.request('GET', url)
    if req.status != 200:
        print("Error:", req.status, "skipping page", url)
        return None
    page = req.data
    soup = BeautifulSoup(page, 'lxml')
    art = soup.find("div", {'class': "post-wrapper"})
    #title and date
    title = (art.find('h1', {'class': 'page-title'})).string
    date = "July 1, 2017"
    post = Post(title, date)
    #text
    body = art.find("div", {"class": 'entry'})
    i = 0
    for paragraph in body.find_all('p'):
        if paragraph.find('em') is not None:
            if i == 0:
                post.set_subtitle((paragraph.find('em')).string)
            else:
                post.set_endnotes((paragraph.find('em')).string)
        else:
            post.add_paragraph(paragraph.string)
        i += 1
    #images
    for i in body.find_all('img'):
        imgUrl = str(i['data-orig-file'])
        imgName = imgUrl.split(".com")[1].replace("/", "_")
        post.add_image([imgName, None]) # name, caption
        #urllib.request.urlretrieve(imgUrl, os.path.basename("img/" + imgName))
    print("done with", url)
    return global_articles[:-17] + [post] + global_articles[-17:]


#processes everything into the word document
def make_doc():
    document = Document()
    document.add_heading('Overneath It All', 0)
    document.add_paragraph("Title Page")
    document.add_page_break()
    document.add_heading("Table of Contents", level=1)
    for art in global_articles:
        document.add_paragraph(art.title + " ... " + art.date)
    document.add_page_break()
    for art in global_articles:
        document.add_heading(art.title, level=1)
        document.add_paragraph(art.date)
        if art.subtitle:
            p = document.add_paragraph()
            p.add_run(art.subtitle).italic = True
        for photo in art.photos:
            try:
                document.add_picture('img/' + photo[0], width=Inches(2.5))
                if photo[1] is not None:
                    document.add_paragraph(photo[1])
            except: #some stuff is tagged as photos but is not
                pass
        for p in art.paragraphs:
            document.add_paragraph(str(p))
        if art.endnotes:
            p = document.add_paragraph()
            p.add_run(art.endnotes).italic = True
        document.add_page_break()
    document.save("overneathitall.docx")


if __name__ == "__main__":
    global_articles = []
    links = make_links()
    for link in links:
        print(link)
        main_scrape("https://overneathitall.com" + link)
    print(len(global_articles))
    global_articles = special_scrape("https://overneathitall.com/living-small/")
    print(len(global_articles))
    make_doc()
    print("done")
